import streamlit as st
from dotenv import load_dotenv
import os
import google.generativeai as genai
from docx import Document
import tempfile

load_dotenv()  # Load environment variables

# Configure Google Gemini API
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Prompt for generating travel itinerary
prompt = """You are a travel itinerary planner. Based on the information provided below:
Departure: {departure}
Destination: {destination}
Days: {days}
Interests: {interests}
Budget: {budget}
Create a detailed travel itinerary that covers important activities, places to visit, and recommendations. Please ensure that the itinerary is well-balanced based on the budget and interests provided.
"""

# Function to generate itinerary using Google Gemini API
def generate_gemini_content(departure, destination, days, interests, budget):
    model = genai.GenerativeModel("gemini-pro")
    full_prompt = prompt.format(departure=departure, destination=destination, days=days, interests=interests, budget=budget)
    response = model.generate_content(full_prompt)
    return response.text

# Function to create and download DOCX
def create_docx(itinerary, departure, destination):
    doc = Document()
    doc.add_heading('Travel Itinerary', 0)
    
    doc.add_paragraph(f"From: {departure} To: {destination}")
    doc.add_paragraph("Itinerary Details:")
    
    doc.add_paragraph(itinerary)

    # Create temporary file to store the document
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
        doc.save(tmp_docx.name)
        return tmp_docx.name

# Custom CSS for styling
st.markdown("""
    <style>
        /* Main container for left alignment */
        .main-container {
            display: flex;
            flex-direction: column;
            align-items: flex-start;
            margin-left: 10%;
            margin-right: 10%;
        }

        /* Styling for the input form section */
        .box {
            width: 100%;
            background-color: #f0f4f8;
            padding: 20px;
            border-radius: 10px;
            box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
            margin-bottom: 20px;
        }

        h1, h2 {
            color: #2C3E50;
            font-family: 'Helvetica', sans-serif;
            font-weight: 700;
        }

        /* Green button */
        .stButton button {
            background-color: #4CAF50;
            color: white;
            border-radius: 5px;
            padding: 10px 20px;
            font-size: 16px;
            margin-top: 20px;
        }

        /* Green button hover effect */
        .stButton button:hover {
            background-color: #45a049;
        }

        /* Result box styling */
        .result-box {
            width: 100%;
            background-color: #ffffff;
            padding: 20px;
            border-radius: 10px;
            box-shadow: 0 2px 6px rgba(0, 0, 0, 0.1);
        }

    </style>
""", unsafe_allow_html=True)

# Streamlit app layout
st.markdown("<div class='main-container'>", unsafe_allow_html=True)

st.title("🌍 Wander Plan")

# Input section wrapped in a styled box
st.markdown("<div class='box'>", unsafe_allow_html=True)
st.header("Travel Plan Details")

# Input fields for the travel itinerary
departure = st.text_input("Enter Departure City:")
destination = st.text_input("Enter Destination City:")
days = st.number_input("Enter Number of Days:", min_value=1, max_value=30, step=1)
interests = st.text_input("Enter Interests (e.g., adventure, culture, food):")
budget = st.selectbox("Select Budget:", options=["Low", "Medium", "High"])

st.markdown("</div>", unsafe_allow_html=True)

# Button to generate the itinerary
if st.button("Generate Itinerary"):
    if departure and destination and days and interests and budget:
        # Generate the itinerary
        itinerary = generate_gemini_content(departure, destination, days, interests, budget)

        # Display the generated itinerary in a styled box below the inputs
        st.markdown("<div class='result-box'>", unsafe_allow_html=True)
        st.header("Generated Travel Itinerary")
        st.write(itinerary)
        st.markdown("</div>", unsafe_allow_html=True)

        # Create a downloadable DOCX file
        docx_file = create_docx(itinerary, departure, destination)

        # Add download button for the DOCX file
        with open(docx_file, "rb") as file:
            st.download_button(
                label="Download Itinerary as DOCX",
                data=file,
                file_name=f"itinerary_{departure}_to_{destination}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.warning("Please fill in all the fields to generate the itinerary.")

st.markdown("</div>", unsafe_allow_html=True)
